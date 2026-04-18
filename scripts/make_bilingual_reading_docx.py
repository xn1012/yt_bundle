from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SECTION_RE = re.compile(r"^##\s+(.+)$")


@dataclass
class Section:
    heading: str
    paragraphs: list[str]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a bilingual reading docx from paired English and Chinese reading markdown files."
    )
    parser.add_argument("english_md", help="Path to the English reading markdown file")
    parser.add_argument("chinese_md", help="Path to the Chinese reading markdown file")
    parser.add_argument(
        "--output",
        help="Output .docx path. Defaults to a same-directory bilingual reading docx beside the English markdown.",
    )
    return parser.parse_args()


def bilingual_docx_output_path(english_md: Path) -> Path:
    stem = english_md.stem
    suffix = " 阅读整理稿"
    if stem.endswith(suffix):
        base_name = stem[: -len(suffix)]
        return english_md.with_name(f"{base_name} 双语阅读整理稿.docx")
    return english_md.with_suffix(".bilingual.docx")


def parse_sections(md_path: Path) -> tuple[str, list[Section]]:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = ""
    sections: list[Section] = []
    current_heading: str | None = None
    current_paragraphs: list[str] = []
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        paragraph = " ".join(part.strip() for part in buffer if part.strip()).strip()
        if paragraph:
            current_paragraphs.append(paragraph)
        buffer = []

    def flush_section() -> None:
        nonlocal current_heading, current_paragraphs
        flush_paragraph()
        if current_heading is not None:
            sections.append(Section(heading=current_heading, paragraphs=current_paragraphs))
        current_heading = None
        current_paragraphs = []

    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if line.startswith("- Source:") or line.startswith("- Generated:") or line.startswith("- Paragraphs:"):
            continue
        match = SECTION_RE.match(line)
        if match:
            flush_section()
            current_heading = match.group(1).strip()
            continue
        if current_heading is None:
            continue
        if not line.strip():
            flush_paragraph()
            continue
        buffer.append(line.strip())

    flush_section()
    return title, sections


def apply_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)


def add_labeled_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label} ")
    label_run.bold = True
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(8)


def add_page_number_field(paragraph) -> None:
    page_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    page_run._r.append(begin)
    page_run._r.append(instr)
    page_run._r.append(separate)
    page_run._r.append(end)


def add_centered_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        add_page_number_field(paragraph)


def write_bilingual_docx(english_md: Path, chinese_md: Path, output_path: Path) -> Path:
    english_title, english_sections = parse_sections(english_md)
    _, chinese_sections = parse_sections(chinese_md)

    document = Document()
    apply_default_font(document)

    title = english_title.replace("Reading Draft", "双语阅读稿").strip()
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph("按 section 对齐排版，英文段落后紧跟对应中文段落。")

    total_sections = max(len(english_sections), len(chinese_sections))
    for index in range(total_sections):
        english_section = english_sections[index] if index < len(english_sections) else None
        chinese_section = chinese_sections[index] if index < len(chinese_sections) else None
        section_heading = (
            english_section.heading
            if english_section is not None
            else chinese_section.heading if chinese_section is not None else f"Section {index + 1}"
        )
        document.add_heading(section_heading, level=1)

        english_paragraphs = english_section.paragraphs if english_section is not None else []
        chinese_paragraphs = chinese_section.paragraphs if chinese_section is not None else []
        total_paragraphs = max(len(english_paragraphs), len(chinese_paragraphs))
        for paragraph_index in range(total_paragraphs):
            if paragraph_index < len(english_paragraphs):
                add_labeled_paragraph(document, "EN", english_paragraphs[paragraph_index])
            if paragraph_index < len(chinese_paragraphs):
                add_labeled_paragraph(document, "ZH", chinese_paragraphs[paragraph_index])

    add_centered_page_numbers(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> int:
    args = parse_args()
    english_md = Path(args.english_md).expanduser().resolve()
    chinese_md = Path(args.chinese_md).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve() if args.output else bilingual_docx_output_path(english_md)
    write_bilingual_docx(english_md=english_md, chinese_md=chinese_md, output_path=output_path)
    print(output_path, flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
